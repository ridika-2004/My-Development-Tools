# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('This is Shiba Inu', 0)

# Image in its native size
doc.add_heading('Image in Native Size:', 3)
doc.add_picture('shiba.png')

# Now save the document to a location
doc.save('shiba.docx')
